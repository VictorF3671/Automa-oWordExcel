import tkinter as tk
from tkinter import filedialog, messagebox
import pandas as pd
from docx import Document
import os

# Função para substituir texto no documento
def substituir_texto(doc, marcador, substituto):
    for paragrafo in doc.paragraphs:
        if marcador in paragrafo.text:
            paragrafo.text = paragrafo.text.replace(marcador, substituto)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                substituir_texto(celula, marcador, substituto)

# Função para carregar a planilha e o modelo
def carregar_planilha_e_modelo():
    global df, caminho_modelo
    try:
        # Selecionar a planilha Excel
        caminho_planilha = filedialog.askopenfilename(filetypes=[("Arquivos Excel", "*.xlsx")])
        df = pd.read_excel(caminho_planilha)
        label_planilha.config(text=f"Planilha selecionada: {caminho_planilha.split('/')[-1]}")

        # Selecionar o modelo de documento Word
        caminho_modelo = filedialog.askopenfilename(filetypes=[("Arquivos Word", "*.docx")])
        label_modelo.config(text=f"Modelo Word selecionado: {caminho_modelo.split('/')[-1]}")

        # Exibir as colunas disponíveis da planilha para o usuário selecionar
        colunas_listbox.delete(0, tk.END)  # Limpar a listbox anterior
        for coluna in df.columns:
            if "Unnamed" not in coluna:
                colunas_listbox.insert(tk.END, coluna)

    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao carregar a planilha ou modelo: {str(e)}")

# Função para selecionar o diretório de saída
def selecionar_diretorio():
    global caminho_diretorio
    caminho_diretorio = filedialog.askdirectory()
    if caminho_diretorio:
        label_diretorio.config(text=f"Diretório selecionado: {caminho_diretorio}")
    else:
        caminho_diretorio = None
        label_diretorio.config(text="Nenhum diretório selecionado")

# Função para gerar os documentos
def gerar_documentos():
    try:
        if not df.empty and caminho_modelo and caminho_diretorio:
            for index, row in df.iterrows():
                doc = Document(caminho_modelo)

                for marcador in df.columns:
                    if marcador in row:
                        substituir_texto(doc, f'{{{{ {marcador} }}}}', str(row[marcador]))

                # Salvar o documento no diretório escolhido pelo usuário
                doc.save(os.path.join(caminho_diretorio, f'TERMO DE DOACAO - {row["nome"]}.docx'))

            messagebox.showinfo("Sucesso", "Documentos gerados com sucesso!")
        else:
            messagebox.showwarning("Aviso", "Verifique se a planilha, o modelo Word e o diretório de destino foram selecionados corretamente.")
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao gerar documentos: {str(e)}")

# Criar a interface gráfica
janela = tk.Tk()
janela.title("Gerador de Documentos Word")

# Definir o tamanho da janela para 1280x720
janela.geometry('1280x600')

# Widgets para exibir o nome dos arquivos selecionados
label_planilha = tk.Label(janela, text="Nenhuma planilha selecionada", anchor='w')
label_planilha.pack(pady=10, fill='x')

label_modelo = tk.Label(janela, text="Nenhum modelo Word selecionado", anchor='w')
label_modelo.pack(pady=10, fill='x')

label_diretorio = tk.Label(janela, text="Nenhum diretório selecionado", anchor='w')
label_diretorio.pack(pady=10, fill='x')

# Botão para carregar os arquivos
btn_carregar = tk.Button(janela, text="Selecionar Planilha e Modelo Word", command=carregar_planilha_e_modelo)
btn_carregar.pack(pady=10)

# Botão para selecionar o diretório de saída
btn_diretorio = tk.Button(janela, text="Selecionar Diretório de Destino", command=selecionar_diretorio)
btn_diretorio.pack(pady=10)

# Listbox para exibir as colunas da planilha
colunas_label = tk.Label(janela, text="Colunas da planilha disponíveis:")
colunas_label.pack(pady=10)

colunas_listbox = tk.Listbox(janela, height=10, width=50)
colunas_listbox.pack(pady=10)

# Botão para gerar os documentos
btn_gerar = tk.Button(janela, text="Gerar Documentos", command=gerar_documentos)
btn_gerar.pack(pady=20)

# Dicionário para armazenar as colunas selecionadas
widgets = {}

janela.mainloop()