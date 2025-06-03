import streamlit as st
import pandas as pd
from docx import Document
import os

def substituir_texto_formatado(paragrafo, marcador, substituto):
    for run in paragrafo.runs:
        if marcador in run.text:
            run.text = run.text.replace(marcador, substituto)

def substituir_texto_no_documento(doc, marcador, substituto):
    for paragrafo in doc.paragraphs:
        substituir_texto_formatado(paragrafo, marcador, substituto)
    
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    substituir_texto_formatado(paragrafo, marcador, substituto)

st.title("Instituto Abilio Pontes")
st.text(" Gere documentos em Word a partir de uma lista do Excel")


caminho_planilha = st.file_uploader("Selecione a planilha Excel:", type=["xlsx"])
df = None
if caminho_planilha:
    df = pd.read_excel(caminho_planilha)
    st.write(f"Planilha selecionada: {caminho_planilha.name}")
    st.write("Prévia da planilha:", df.head())


caminho_modelo = st.file_uploader("Selecione o modelo do documento Word:", type=["docx"])
if caminho_modelo:
    st.write(f"Modelo Word selecionado: {caminho_modelo.name}")

caminho_diretorio = st.text_input("Caminho do diretório para salvar os documentos:", "")
if caminho_diretorio and not os.path.isdir(caminho_diretorio):
    st.warning("Por favor, insira um caminho de diretório válido.")


padrao_nome = st.text_input("Formato do nome do arquivo (ex: 'TERMOS - {nome}'):", "CERTIFICADO - {nome}")


if st.button("Gerar Documentos"):
    if df is not None and caminho_modelo and caminho_diretorio:
        for index, row in df.iterrows():
            doc = Document(caminho_modelo)
            
            for coluna in df.columns:
                if pd.notna(row[coluna]):
                    substituir_texto_no_documento(doc, f'{{{{ {coluna} }}}}', str(row[coluna]))
            
           
            nome_arquivo = padrao_nome
            for coluna in df.columns:
                if f'{{{coluna}}}' in nome_arquivo:
                    nome_arquivo = nome_arquivo.replace(f'{{{coluna}}}', str(row[coluna]))
            
            nome_arquivo += ".docx"
            
            doc.save(os.path.join(caminho_diretorio, nome_arquivo))
        
        st.success("Documentos gerados com sucesso!")
    else:
        st.error("Por favor, selecione a planilha, o modelo Word e o diretório de destino.")