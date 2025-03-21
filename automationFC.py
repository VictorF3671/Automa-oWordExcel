import pandas as pd
from docx import Document

# Função para substituir texto no documento
def substituir_texto(doc, marcador, substituto):
    for paragrafo in doc.paragraphs:
        if marcador in paragrafo.text:
            paragrafo.text = paragrafo.text.replace(marcador, substituto)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                substituir_texto(celula, marcador, substituto)

# Carregar dados da planilha Excel
df = pd.read_excel('lista.xlsx')
count = 1
# Iterar sobre cada linha da planilha e gerar um documento Word para cada uma
for index, row in df.iterrows():
    # Carregar o documento modelo
    doc = Document('modeloFC.docx')
    
    # Substituir os marcadores de posição pelo valor correspondente
    substituir_texto(doc, '{{ nome }}', row['nome'])
    substituir_texto(doc, '{{ cpf }}', row['cpf'])
    substituir_texto(doc, '{{ rg }}', row['rg'])
    substituir_texto(doc, '{{ endereco }}', row['endereço'])
    substituir_texto(doc, '{{ bairro }}', row['bairro'])
    substituir_texto(doc, '{{ cep }}', row['cep'])
    substituir_texto(doc, '{{ telefone }}', row['telefone'])
     
    substituir_texto(doc, '{{ renda }}', str(row['renda'])) 

    # Obter a data de nascimento
    data_nascimento = row['dNasc']

    # Verificar se o valor é do tipo datetime e formatá-lo
    if isinstance(data_nascimento, pd.Timestamp):
        data_formatada = data_nascimento.strftime('%d/%m/%Y')
    else:
        data_formatada = str(data_nascimento)
    substituir_texto(doc, '{{ dNasc }}', data_formatada)

    if row['tCasa'] == 'PROPRIA' or row['tCasa'] == 'PRÓPRIA' :
        substituir_texto(doc, '{{ pr }}', 'X')
    else : 
        substituir_texto(doc, '{{ pr }}', ' ')

    if row['tCasa'] == 'ALUGADA' :
        substituir_texto(doc, '{{ al }}', 'X')
    else : 
        substituir_texto(doc, '{{ al }}', ' ')
        
    if row['tCasa'] == 'CEDIDA' :
        substituir_texto(doc, '{{ ced }}', 'X')
    else : 
        substituir_texto(doc, '{{ ced }}', '')

    if row['estadoC'] == 'SOLTEIRO' or row['estadoC'] == 'SOLTEIRA':
        substituir_texto(doc, '{{ sol }}', 'X')
    else : 
        substituir_texto(doc, '{{ sol }}', '')

    if row['estadoC'] == 'CASADO' or row['estadoC'] == 'CASADA':
        substituir_texto(doc, '{{ cas }}', 'X')
    else : 
        substituir_texto(doc, '{{ cas }}', '')

    if row['estadoC'] == 'DIVORCIADO' or row['estadoC'] == 'DIVORCIADA' or row['estadoC'] == 'SEPARADO' or row['estadoC'] == 'SEPARADA':
        substituir_texto(doc, '{{ div }}', 'X')
    else : 
        substituir_texto(doc, '{{ div }}', '')
    
    if row['estadoC'] == 'UNIAO ESTAVEL' or row['estadoC'] == 'UNIÃO ESTÁVEL' or row['estadoC'] == 'UNIAO ESTÁVEL' or row['estadoC'] == 'UNIÃO ESTAVEL':
        substituir_texto(doc, '{{ ue }}', 'X')
    else : 
        substituir_texto(doc, '{{ ue }}', '')

    if row['estadoC'] == 'VIÚVA' or row['estadoC'] == 'VIÚVO' or row['estadoC'] == 'VIUVA' or row['estadoC'] == 'VIUVO':
        substituir_texto(doc, '{{ viu }}', 'X')
    else : 
        substituir_texto(doc, '{{ viu }}', '')
    
    if row['escolaridade'] == 'ENSINO FUNDAMENTAL COMPLETO' or row['escolaridade'] == 'ENS. FUNDAMENTAL COMPLETO' or row['escolaridade'] == 'ALFABETIZADA' or row['escolaridade'] == 'ENS. FUNDAMENTAL INCOMPLETO' or row['escolaridade'] == 'ENSINO FUNDAMENTAL INCOMPLETO' or row['escolaridade'] == 'ENS.FUNDAMENTAL COMPLETO' or row['escolaridade'] == 'ENS.FUNDAMENTAL INCOMPLETO':
        substituir_texto(doc, '{{ fund }}', 'X')
    else : 
        substituir_texto(doc, '{{ fund }}', '')

    if row['escolaridade'] == 'ENSINO MEDIO COMPLETO' or row['escolaridade'] == 'ENS. MEDIO COMPLETO' or row['escolaridade'] == 'ENS. MÉDIO INCOMPLETO' or row['escolaridade'] == 'ENSINO MÉDIO INCOMPLETO' or row['escolaridade'] == 'ENS.MÉDIO INCOMPLETO' or row['escolaridade'] == 'ENS. MEDIO INCOMPLETO' or  row['escolaridade'] == 'ENS. MÉDIO INCOMPLETO' :
        substituir_texto(doc, '{{ med }}', 'X')
    else : 
        substituir_texto(doc, '{{ med }}', '')

    if row['profissao'] == 'DO LAR' or row['profissao'] == 'SEM RENDA' or row['profissao'] == 'SEM OCUPAÇÃO' or row['profissao'] == 'APOSENTADO' or row['profissao'] == 'APOSENTADA' or row['profissao'] == 'NENHUM' or row['profissao'] == 'SEM REGISTRO' or row['profissao'] == 'B.P.C'or row['profissao'] == 'PENSIONISTA' or row['profissao'] == 'PENSÃO':
        substituir_texto(doc, '{{ tn }}', 'X')
        substituir_texto(doc, '{{ ts }}', ' ')
        substituir_texto(doc, ' {{ profi }}', '')
    else : 
        substituir_texto(doc, '{{ tn }}', '')
        substituir_texto(doc, ' {{ ts }}', 'X')
        substituir_texto(doc, ' {{ profi }}', row['profissao'])

    if row['profissao'] == 'BPC' or row['profissao'] == 'B.P.C' :
        substituir_texto(doc, '{{ bpc }}', 'X')
    else : 
        substituir_texto(doc, '{{ bpc }}', '')

    if row['profissao'] == 'BOLSA FAMILIA' or row['profissao'] == 'BOLSA FAMÍLIA' :
        substituir_texto(doc, '{{ bf }}', 'X')
    else : 
        substituir_texto(doc, '{{ bf }}', '')

    if row['profissao'] == 'APOSENTADO' or row['profissao'] == 'APOSENTADA'or row['profissao'] == 'PENSIONISTA' or row['profissao'] == 'PENSÃO' or row['profissao'] == 'AUX. ESTADUAL' or row['profissao'] == 'AUXILIO ESTADUAL' :
        substituir_texto(doc, '{{ otr }}', 'X')
    else : 
        substituir_texto(doc, '{{ otr }}', '')

    substituir_texto(doc, '{{ count }}', str(count))    

    count += 1
    # Salvar o documento resultante
    doc.save(f'FICHA DE CADASTRO - {row['nome']}.docx')

print("Documentos gerados com sucesso!")